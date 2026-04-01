"""
Document loaders — one adapter per source type.
Each loader returns a list of Document objects with extracted text + metadata.
"""
from __future__ import annotations
import logging
import time
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument

from utils.models import Document, DocumentMetadata

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Base interface
# ─────────────────────────────────────────────────────────────────────────────

class BaseLoader(ABC):
    """All loaders implement this interface."""

    @abstractmethod
    def load(self, source: str | bytes | Path, **kwargs) -> list[Document]:
        """Load one or more documents from the given source."""
        ...

    @abstractmethod
    def supports(self, source: str | Path) -> bool:
        """Return True if this loader can handle the given source."""
        ...


# ─────────────────────────────────────────────────────────────────────────────
# PDF loader
# ─────────────────────────────────────────────────────────────────────────────

class PDFLoader(BaseLoader):
    """
    Loads PDFs using pdfplumber for digital PDFs.
    """

    def supports(self, source: str | Path) -> bool:
        return str(source).lower().endswith(".pdf")

    def load(self, source: str | bytes | Path, **kwargs) -> list[Document]:
        start_time = time.time()
        filename = str(source) if isinstance(source, Path) else source

        with pdfplumber.open(source) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            page_count = len(pdf.pages)
            word_count = len(text.split())
            load_time = time.time() - start_time

            metadata = DocumentMetadata(
                source=filename,
                doc_type="pdf",
                page_count=page_count,
                custom={"word_count": word_count, "load_time": load_time}
            )

            return [Document(content=text.strip(), metadata=metadata)]


# ─────────────────────────────────────────────────────────────────────────────
# DOCX loader
# ─────────────────────────────────────────────────────────────────────────────

class DOCXLoader(BaseLoader):
    """
    Loads Microsoft Word documents.
    """

    def supports(self, source: str | Path) -> bool:
        return str(source).lower().endswith(".docx")

    def load(self, source: str | bytes | Path, **kwargs) -> list[Document]:
        start_time = time.time()
        filename = str(source) if isinstance(source, Path) else source

        doc = DocxDocument(source)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"

        page_count = None  # DOCX doesn't have pages easily
        word_count = len(text.split())
        load_time = time.time() - start_time

        metadata = DocumentMetadata(
            source=filename,
            doc_type="docx",
            page_count=page_count,
            custom={"word_count": word_count, "load_time": load_time}
        )

        return [Document(content=text.strip(), metadata=metadata)]


# ─────────────────────────────────────────────────────────────────────────────
# TXT loader
# ─────────────────────────────────────────────────────────────────────────────

class TXTLoader(BaseLoader):
    """
    Loads plain text files.
    """

    def supports(self, source: str | Path) -> bool:
        return str(source).lower().endswith(".txt")

    def load(self, source: str | bytes | Path, **kwargs) -> list[Document]:
        start_time = time.time()
        filename = str(source) if isinstance(source, Path) else source

        with open(source, 'r', encoding='utf-8') as f:
            text = f.read()

        page_count = None
        word_count = len(text.split())
        load_time = time.time() - start_time

        metadata = DocumentMetadata(
            source=filename,
            doc_type="txt",
            page_count=page_count,
            custom={"word_count": word_count, "load_time": load_time}
        )

        return [Document(content=text.strip(), metadata=metadata)]


# ─────────────────────────────────────────────────────────────────────────────
# CSV loader
# ─────────────────────────────────────────────────────────────────────────────

class CSVLoader(BaseLoader):
    """
    Loads CSV files as text.
    """

    def supports(self, source: str | Path) -> bool:
        return str(source).lower().endswith(".csv")

    def load(self, source: str | bytes | Path, **kwargs) -> list[Document]:
        start_time = time.time()
        filename = str(source) if isinstance(source, Path) else source

        df = pd.read_csv(source)
        text = df.to_string()

        page_count = None
        word_count = len(text.split())
        load_time = time.time() - start_time

        metadata = DocumentMetadata(
            source=filename,
            doc_type="csv",
            page_count=page_count,
            custom={"word_count": word_count, "load_time": load_time}
        )

        return [Document(content=text.strip(), metadata=metadata)]


# ─────────────────────────────────────────────────────────────────────────────
# Loader registry
# ─────────────────────────────────────────────────────────────────────────────

loader_registry = [
    PDFLoader(),
    DOCXLoader(),
    TXTLoader(),
    CSVLoader(),
]


def get_loader(source: str | Path) -> Optional[BaseLoader]:
    """Get the appropriate loader for the source."""
    for loader in loader_registry:
        if loader.supports(source):
            return loader
    return None

    def load(self, source, **kwargs) -> list[Document]:
        import pdfplumber  # pip install pdfplumber

        path = Path(source)
        pages_text: list[str] = []
        page_count = 0

        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                if not text.strip():
                    text = self._ocr_page(page)
                pages_text.append(text)

        # Combine all pages into a single document with page markers
        full_text = "\n\n".join(
            f"[Page {i+1}]\n{t}" for i, t in enumerate(pages_text) if t.strip()
        )

        metadata = DocumentMetadata(
            source=str(path),
            doc_type="pdf",
            title=kwargs.get("title") or path.stem,
            page_count=page_count,
        )
        return [Document(content=full_text, metadata=metadata)]

    def _ocr_page(self, page) -> str:
        try:
            import pytesseract  # pip install pytesseract

            img = page.to_image(resolution=200).original
            return pytesseract.image_to_string(img)
        except ImportError:
            logger.warning("pytesseract not installed; skipping OCR for scanned page.")
            return ""


# ─────────────────────────────────────────────────────────────────────────────
# Plain text / Markdown loader
# ─────────────────────────────────────────────────────────────────────────────

class TextLoader(BaseLoader):
    EXTENSIONS = {".txt", ".md", ".rst", ".log"}

    def supports(self, source) -> bool:
        return Path(str(source)).suffix.lower() in self.EXTENSIONS

    def load(self, source, encoding="utf-8", **kwargs) -> list[Document]:
        path = Path(source)
        content = path.read_text(encoding=encoding)
        metadata = DocumentMetadata(
            source=str(path),
            doc_type=path.suffix.lstrip("."),
            title=kwargs.get("title") or path.stem,
        )
        return [Document(content=content, metadata=metadata)]


# ─────────────────────────────────────────────────────────────────────────────
# DOCX loader
# ─────────────────────────────────────────────────────────────────────────────

class DocxLoader(BaseLoader):
    def supports(self, source) -> bool:
        return str(source).lower().endswith(".docx")

    def load(self, source, **kwargs) -> list[Document]:
        from docx import Document as DocxDocument  # pip install python-docx

        path = Path(source)
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        # Extract tables as TSV blocks
        for table in doc.tables:
            rows = ["\t".join(cell.text for cell in row.cells) for row in table.rows]
            content += "\n\n[TABLE]\n" + "\n".join(rows)

        metadata = DocumentMetadata(
            source=str(path),
            doc_type="docx",
            title=kwargs.get("title") or path.stem,
        )
        return [Document(content=content, metadata=metadata)]


# ─────────────────────────────────────────────────────────────────────────────
# HTML loader
# ─────────────────────────────────────────────────────────────────────────────

class HTMLLoader(BaseLoader):
    def supports(self, source) -> bool:
        s = str(source).lower()
        return s.endswith(".html") or s.endswith(".htm") or s.startswith("http")

    def load(self, source, **kwargs) -> list[Document]:
        from bs4 import BeautifulSoup  # pip install beautifulsoup4

        if str(source).startswith("http"):
            import httpx
            response = httpx.get(str(source), timeout=30, follow_redirects=True)
            html = response.text
            source_str = str(source)
        else:
            html = Path(source).read_text()
            source_str = str(source)

        soup = BeautifulSoup(html, "html.parser")
        # Remove scripts, styles, nav
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        title = soup.title.string.strip() if soup.title else None
        content = soup.get_text(separator="\n", strip=True)

        metadata = DocumentMetadata(
            source=source_str,
            doc_type="html",
            title=title or kwargs.get("title"),
        )
        return [Document(content=content, metadata=metadata)]


# ─────────────────────────────────────────────────────────────────────────────
# CSV loader — treats each row as a small document for structured retrieval
# ─────────────────────────────────────────────────────────────────────────────

class CSVLoader(BaseLoader):
    def supports(self, source) -> bool:
        return str(source).lower().endswith(".csv")

    def load(self, source, row_as_doc: bool = False, **kwargs) -> list[Document]:
        import csv
        path = Path(source)
        documents: list[Document] = []

        with open(path, newline="", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            rows = list(reader)

        if row_as_doc:
            # Each row becomes its own document (good for structured records)
            for i, row in enumerate(rows):
                content = " | ".join(f"{k}: {v}" for k, v in row.items() if v.strip())
                meta = DocumentMetadata(
                    source=str(path),
                    doc_type="csv",
                    title=f"{path.stem} row {i}",
                    custom=dict(row),
                )
                documents.append(Document(content=content, metadata=meta))
        else:
            # Whole CSV as single document
            content = "\n".join(
                " | ".join(f"{k}: {v}" for k, v in row.items()) for row in rows
            )
            meta = DocumentMetadata(source=str(path), doc_type="csv", title=path.stem)
            documents.append(Document(content=content, metadata=meta))

        return documents


# ─────────────────────────────────────────────────────────────────────────────
# Loader registry
# ─────────────────────────────────────────────────────────────────────────────

class LoaderRegistry:
    """Dispatcher: given a source path, pick the right loader automatically."""

    def __init__(self):
        self._loaders: list[BaseLoader] = [
            PDFLoader(),
            DocxLoader(),
            HTMLLoader(),
            CSVLoader(),
            TextLoader(),  # fallback last
        ]

    def get_loader(self, source: str | Path) -> BaseLoader:
        for loader in self._loaders:
            if loader.supports(source):
                return loader
        raise ValueError(f"No loader found for source: {source}")

    def load(self, source: str | Path, **kwargs) -> list[Document]:
        loader = self.get_loader(source)
        logger.info("Loading %s with %s", source, type(loader).__name__)
        return loader.load(source, **kwargs)


loader_registry = LoaderRegistry()
